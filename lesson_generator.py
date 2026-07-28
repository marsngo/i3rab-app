import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_rtl(paragraph):
    """تحديد اتجاه الفقرة من اليمين إلى اليسار RTL"""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def create_formatted_lesson_docx(lesson_title: str, lesson_content: str, output_path: str = "lesson_plan.docx") -> str:
    """
    دالة تقوم بإنشاء ملف Word تنسيقي احترافي لخطة الدرس باللغة العربية
    """
    doc = Document()
    
    # ضبط الهوامش
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. عنوان الدرس الرئيسية
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(title_para)
    run = title_para.add_run(f"خطة درس: {lesson_title}")
    run.font.name = 'Traditional Arabic'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 54, 93) # أزرق كحلي فخم

    # خط فاصل بسيط
    doc.add_paragraph()

    # 2. تقسيم محتوى الدرس وإضافته
    lines = lesson_content.split('\n')
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # العناوين الرئيسية
        if line_str.startswith("###") or line_str.startswith("##") or line_str.startswith("#"):
            clean_title = line_str.replace("#", "").strip()
            run = p.add_run(clean_title)
            run.font.name = 'Traditional Arabic'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(140, 29, 64) # عنابي أنيق
        # الفقرات العادية والقوائم
        else:
            run = p.add_run(line_str)
            run.font.name = 'Traditional Arabic'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(30, 30, 30)

    # حفظ الملف
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    # تجربة سريعة لتوليد ملف ورد تجريبي
    sample_content = """
    ## الأهداف التعليمية
    - أن يتعرف الطالب على مفاهيم الفاعل وأحكامه النحوية.
    - أن يستخرج الطالب الفاعل من النص ويضبطه بالشكل.
    
    ## الشرح والقاعدة
    الفاعل هو اسم مرفوع أُسند إليه فعل تام متقدم عليه.
    
    ## أمثلة وتطبيقات
    - جاءَ القَاضِي إلى المَحْكَمَةِ.
    - ﴿وَجَاءَ رَجُلٌ مِّنْ أَقْصَى الْمَدِينَةِ يَسْعَى﴾ [سورة قصص: 20]
    """
    created_file = create_formatted_lesson_docx("الفاعل وأحكامه", sample_content, "test_lesson.docx")
    print(f"تم إنشاء ملف الدرس بنجاح: {created_file}")